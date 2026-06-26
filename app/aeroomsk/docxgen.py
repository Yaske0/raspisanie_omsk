# -*- coding: utf-8 -*-
"""
Генерация документа Word по образцу пользователя.

Шрифт Verdana: заголовки 28, текст 11. Две секции — «Прилет» и «Вылет»,
каждая таблицей с рамками. Поля узкие, чтобы уместить на одну страницу.
"""

from __future__ import annotations

import datetime as dt

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .builder import Row

FONT = "Verdana"
TEXT_PT = 11
HEAD_PT = 28

# доли ширины колонок: Дата, Город, Рейс, Авиакомпания, Время, Статус
COL_WEIGHTS = [0.086, 0.294, 0.108, 0.269, 0.092, 0.151]


def _set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tcPr.append(borders)


def _set_cell_margins(cell, top=8, bottom=8, left=70, right=70):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        m.append(el)
    tcPr.append(m)


def _write_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text or "")
    run.font.name = FONT
    run.font.size = Pt(TEXT_PT)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:cs"), FONT)
    _set_cell_border(cell)
    _set_cell_margins(cell)


def _add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(HEAD_PT)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn("w:cs"), FONT)


def _set_fixed_layout(table, widths):
    tbl = table._tbl
    tblPr = tbl.tblPr
    # фиксированная раскладка — иначе Word подгоняет ширину под текст и переносит
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    total = OxmlElement("w:tblW")
    total.set(qn("w:w"), str(int(sum(w.twips for w in widths))))
    total.set(qn("w:type"), "dxa")
    tblPr.append(total)
    # явная сетка столбцов
    grid = tbl.tblGrid
    for col in list(grid):
        grid.remove(col)
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w.twips)))
        grid.append(gc)


def _add_table(doc, rows: list[Row], content_width_cm: float):
    table = doc.add_table(rows=len(rows), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False
    widths = [Cm(content_width_cm * w) for w in COL_WEIGHTS]
    _set_fixed_layout(table, widths)
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i].cells
        values = [row.date, row.city, row.flight_no, row.airline, row.sched, row.status]
        for c_i, val in enumerate(values):
            _write_cell(cells[c_i], val)
            cells[c_i].width = widths[c_i]
    # высота строк — по содержимому, без лишнего запаса
    for r in table.rows:
        trPr = r._tr.get_or_add_trPr()
        h = OxmlElement("w:trHeight")
        h.set(qn("w:val"), "0")
        h.set(qn("w:hRule"), "atLeast")
        trPr.append(h)
    return table


def build_document(arr_rows: list[Row], dep_rows: list[Row],
                   out_path: str, caption: str | None = None) -> str:
    doc = Document()

    # шрифт по умолчанию
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(TEXT_PT)
    normal.element.rPr.rFonts.set(qn("w:cs"), FONT)

    sec = doc.sections[0]
    sec.top_margin = Cm(1.2)
    sec.bottom_margin = Cm(1.0)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    content_w = (sec.page_width - sec.left_margin - sec.right_margin) / 360000  # EMU->cm

    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(caption)
        run.font.name = FONT
        run.font.size = Pt(TEXT_PT)
        run.font.italic = True

    _add_heading(doc, "Прилет")
    _add_table(doc, arr_rows, content_w)
    _add_heading(doc, "Вылет")
    _add_table(doc, dep_rows, content_w)

    doc.save(out_path)
    return out_path
